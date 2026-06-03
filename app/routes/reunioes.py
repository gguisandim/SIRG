# app/routes/reunioes.py
from flask import Blueprint, redirect, render_template, request, url_for, send_file, jsonify, flash
from datetime import datetime
from flask_login import login_required, current_user
from app import db
from app.models import Reuniao, Frequencia, Membro
from app.services.pauta_reuniao import gerar_pdf_pauta_reuniao
from app.services.ia.groq_client import GroqClient
from app.services.ia.prompts import (
    carregar_modelo_pauta_colegiado,
    prompt_pauta_colegiado_com_modelo,
    carregar_modelo_ata_colegiado,
    prompt_ata_colegiado_com_modelo
)
from app.utils.permissoes import administrativo_required
from flask import Blueprint, redirect, render_template, request, url_for, send_file, jsonify, flash, current_app
from flask_mail import Message
from app import mail
import textwrap
from flask import send_file
from io import BytesIO
from docx import Document

reunioes = Blueprint('reunioes', __name__)


@reunioes.route('/')
@login_required
@administrativo_required
def listar_reunioes():
    lista_reunioes = Reuniao.query.all()
    return render_template('reunioes/listar.html', reunioes=lista_reunioes)


@reunioes.route('/adicionar', methods=['GET', 'POST'])
@login_required
@administrativo_required
def adicionar_reuniao():

    if request.method == 'POST':

        titulo = request.form['titulo']
        data = datetime.strptime(request.form['data'], '%Y-%m-%d').date()
        horario = datetime.strptime(request.form['horario'], '%H:%M').time()
        local = request.form['local']
        tipo = request.form['tipo']
        pauta = request.form.get('pauta')
        numero_oficio = request.form.get('numero_oficio') or None

        reuniao = Reuniao(
            titulo=titulo,
            data=data,
            horario=horario,
            local=local,
            tipo=tipo,
            pauta=pauta,
            numero_oficio=numero_oficio
        )

        db.session.add(reuniao)
        db.session.commit()

        return redirect(url_for('reunioes.listar_reunioes'))

    return render_template('reunioes/adicionar.html')


@reunioes.route('/detalhes/<int:id>', methods=['GET', 'POST'])
@login_required
@administrativo_required
def detalhes_reuniao(id):
    reuniao = Reuniao.query.get_or_404(id)

    professores = Membro.query.filter(
    Membro.ativo == True,
    Membro.tipo.in_(["professor", "tecnico"])
).order_by(Membro.nome).all()

    # Garante que todos os professores tenham registro de frequência nessa reunião
    for professor in professores:
        frequencia_existente = Frequencia.query.filter_by(
            membro_id=professor.id,
            reuniao_id=reuniao.id
        ).first()

        if frequencia_existente is None:
            frequencia = Frequencia(
                membro_id=professor.id,
                reuniao_id=reuniao.id,
                status="falta"
            )
            db.session.add(frequencia)

    db.session.commit()

    if request.method == 'POST':
        for professor in professores:
            status = request.form.get(f'status_{professor.id}')

            if not status:
                continue

            frequencia = Frequencia.query.filter_by(
                membro_id=professor.id,
                reuniao_id=reuniao.id
            ).first()

            if frequencia is None:
                frequencia = Frequencia(
                    membro_id=professor.id,
                    reuniao_id=reuniao.id,
                    status=status
                )
                db.session.add(frequencia)
            else:
                frequencia.status = status

        db.session.commit()

        return redirect(url_for('reunioes.detalhes_reuniao', id=reuniao.id))

    frequencias = Frequencia.query.filter_by(
        reuniao_id=reuniao.id
    ).all()

    frequencias_por_membro = {
        frequencia.membro_id: frequencia
        for frequencia in frequencias
    }

    justificativas = [
        frequencia for frequencia in frequencias
        if frequencia.justificativa
    ]

    return render_template(
        'reunioes/detalhes.html',
        reuniao=reuniao,
        professores=professores,
        frequencias_por_membro=frequencias_por_membro,
        justificativas=justificativas,
        gerado_em=datetime.now().strftime('%d/%m/%Y às %H:%M')
    )

@reunioes.route('/pauta/<int:id>')
@login_required
@administrativo_required
def gerar_pauta_reuniao(id):
    reuniao = Reuniao.query.get_or_404(id)

    caminho_pdf = gerar_pdf_pauta_reuniao(reuniao)

    return send_file(
        caminho_pdf,
        as_attachment=True,
        download_name=f'pauta_reuniao_{reuniao.id}.pdf',
        mimetype='application/pdf'
    )


@reunioes.route('/editar/<int:id>', methods=['GET', 'POST'])
@login_required
@administrativo_required
def editar_reuniao(id):

    reuniao = Reuniao.query.get_or_404(id)

    if request.method == 'POST':

        reuniao.titulo = request.form['titulo']
        reuniao.data = datetime.strptime(request.form['data'], '%Y-%m-%d').date()
        reuniao.horario = datetime.strptime(request.form['horario'], '%H:%M').time()
        reuniao.local = request.form['local']
        reuniao.tipo = request.form['tipo']
        reuniao.numero_oficio = request.form.get('numero_oficio') or None
        reuniao.pauta = request.form.get('pauta')

        db.session.commit()

        return redirect(url_for('reunioes.detalhes_reuniao', id=reuniao.id))

    return render_template('reunioes/editar.html', reuniao=reuniao)


@reunioes.route('/editar/<int:id>/gerar-pauta-ia', methods=['POST'])
@login_required
@administrativo_required
def gerar_pauta_ia_edicao(id):
    reuniao = Reuniao.query.get_or_404(id)

    try:
        dados = request.get_json(silent=True) or {}

        titulo = dados.get('titulo') or reuniao.titulo or ''
        data = dados.get('data') or reuniao.data.strftime('%d/%m/%Y')
        horario = dados.get('horario') or reuniao.horario.strftime('%H:%M')
        local = dados.get('local') or reuniao.local or ''
        tipo = dados.get('tipo') or reuniao.tipo or ''
        pauta_atual = dados.get('pauta') or reuniao.pauta or ''

        modelo_pauta = carregar_modelo_pauta_colegiado()

        prompt = prompt_pauta_colegiado_com_modelo(
            titulo=titulo,
            data_reuniao=data,
            horario=horario,
            local=local,
            tipo=tipo,
            pauta_bruta=pauta_atual,
            modelo_pauta=modelo_pauta
        )

        groq = GroqClient()
        texto = groq.gerar_texto(prompt)

        return jsonify({
            'sucesso': True,
            'texto': texto
        })

    except Exception as erro:
        current_app.logger.exception("Erro ao gerar pauta com IA")

        return jsonify({
            'sucesso': False,
            'erro': str(erro)
        }), 500
    
@reunioes.route('/adicionar/gerar-pauta-ia', methods=['POST'])
@login_required
@administrativo_required
def gerar_pauta_ia_cadastro():
    try:
        dados = request.get_json(silent=True) or {}

        titulo = dados.get('titulo') or ''
        data = dados.get('data') or ''
        horario = dados.get('horario') or ''
        local = dados.get('local') or ''
        tipo = dados.get('tipo') or ''
        pauta_atual = dados.get('pauta') or ''

        modelo_pauta = carregar_modelo_pauta_colegiado()

        prompt = prompt_pauta_colegiado_com_modelo(
            titulo=titulo,
            data_reuniao=data,
            horario=horario,
            local=local,
            tipo=tipo,
            pauta_bruta=pauta_atual,
            modelo_pauta=modelo_pauta
        )

        groq = GroqClient()
        texto = groq.gerar_texto(prompt)

        return jsonify({
            'sucesso': True,
            'texto': texto
        })

    except Exception as erro:
        return jsonify({
            'sucesso': False,
            'erro': str(erro)
        }), 500

@reunioes.route("/<int:id>/pauta/word")
@login_required
@administrativo_required
def baixar_pauta_word(id):
    from io import BytesIO
    from flask import send_file
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    reuniao = Reuniao.query.get_or_404(id)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Pt(56)
    section.bottom_margin = Pt(56)
    section.left_margin = Pt(70)
    section.right_margin = Pt(70)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    cabecalho = doc.add_paragraph()
    cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cabecalho.add_run(
        "UNIVERSIDADE FEDERAL DO PARÁ\n"
        "NÚCLEO DE ESTUDOS TRANSDISCIPLINARES EM EDUCAÇÃO BÁSICA\n"
        "PROGRAMA DE PÓS-GRADUAÇÃO EM CURRÍCULO E GESTÃO DA ESCOLA BÁSICA"
    )
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)

    doc.add_paragraph("")

    linha = doc.add_paragraph()
    linha.add_run(
        f"OFÍCIO Nº {reuniao.numero_oficio or 'None'}/{reuniao.data.year} - PPEB/NEB"
    )
    linha.add_run(
        f"\t\t\t\tBelém, {reuniao.data.strftime('%d/%m/%Y')}."
    )

    doc.add_paragraph("")
    doc.add_paragraph("Da: Coordenação do PPEB")
    doc.add_paragraph("Para: Membros do Colegiado do PPEB")

    doc.add_paragraph("")
    assunto = doc.add_paragraph()
    assunto.add_run("Assunto: ").bold = True
    assunto.add_run(f"{reuniao.titulo} – Convocatória")

    doc.add_paragraph("")
    doc.add_paragraph("Prezados(as) Membros")

    doc.add_paragraph("")
    texto = doc.add_paragraph()
    texto.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    texto.add_run(
        f"Convocamos a todos/as para {reuniao.titulo} a ser realizada no "
    )
    texto.add_run(f"dia {reuniao.data.strftime('%d/%m/%Y')}").bold = True
    texto.add_run(", às ")
    texto.add_run(f"{reuniao.horario.strftime('%Hh')}").bold = True
    texto.add_run(
        f", em {reuniao.local}, para tratar da seguinte pauta:"
    )

    doc.add_paragraph("")

    for item in (reuniao.pauta or "").splitlines():
        linha_pauta = item.strip()
        if not linha_pauta:
            continue

        p = doc.add_paragraph()

        if linha_pauta[:2] in ["1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9."]:
            if linha_pauta.startswith(("1.1", "1.2", "2.1", "2.2", "3.1", "3.2", "3.3", "3.4", "3.5", "4.1", "4.2")):
                p.paragraph_format.left_indent = Pt(24)
                p.add_run(linha_pauta)
            else:
                p.add_run(linha_pauta).bold = True
        else:
            p.paragraph_format.left_indent = Pt(24)
            p.add_run(linha_pauta)

    doc.add_paragraph("")
    doc.add_paragraph("Atenciosamente,")

    doc.add_paragraph("")
    doc.add_paragraph("")

    assinatura = doc.add_paragraph()
    assinatura.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = assinatura.add_run("Émina Santos\n")
    r.bold = True
    r.italic = True
    assinatura.add_run("Coordenadora do PPEB\n")
    assinatura.add_run("Portaria 1342/2026")

    arquivo = BytesIO()
    doc.save(arquivo)
    arquivo.seek(0)

    return send_file(
        arquivo,
        as_attachment=True,
        download_name=f"pauta_reuniao_{reuniao.id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


@reunioes.route("/<int:id>/gerar-ata-ia", methods=["POST"])
@login_required
@administrativo_required
def gerar_ata_ia(id):
    reuniao = Reuniao.query.get_or_404(id)

    try:
        dados = request.get_json(silent=True) or {}

        presentes = (dados.get("presentes") or "").strip()
        justificativas_texto = (dados.get("justificativas") or "").strip()
        anotacoes = (dados.get("anotacoes") or "").strip()
        pauta = (dados.get("pauta") or reuniao.pauta or "").strip()

        if not presentes:
            presentes = "[Nomes dos presentes não informados]"

        if not justificativas_texto:
            justificativas_texto = "[Ausências justificadas não informadas]"

        modelo_ata = carregar_modelo_ata_colegiado()

        prompt = prompt_ata_colegiado_com_modelo(
            titulo=reuniao.titulo or "",
            data_reuniao=reuniao.data.strftime("%d/%m/%Y"),
            horario=reuniao.horario.strftime("%H:%M"),
            local=reuniao.local or "",
            tipo=reuniao.tipo or "",
            pauta=pauta,
            presentes=presentes,
            justificativas=justificativas_texto,
            anotacoes=anotacoes,
            modelo_ata=modelo_ata
        )

        prompt += f"""

REGRA FINAL OBRIGATÓRIA:
Na seção de presentes, escreva exatamente estes nomes:
{presentes}

Na seção de ausências justificadas, escreva exatamente estas informações:
{justificativas_texto}

É proibido substituir os presentes por "[Nomes dos presentes não informados]" quando houver nomes acima.
"""

        groq = GroqClient()
        texto = groq.gerar_texto(prompt)

        return jsonify({
            "sucesso": True,
            "texto": texto
        })

    except Exception as erro:
        current_app.logger.exception("Erro ao gerar ata com IA")

        return jsonify({
            "sucesso": False,
            "erro": str(erro)
        }), 500

@reunioes.route("/<int:id>/ata", methods=["POST"])
@login_required
@administrativo_required
def salvar_ata(id):
    reuniao = Reuniao.query.get_or_404(id)

    reuniao.ata = request.form.get("ata")
    db.session.commit()

    flash("Ata salva com sucesso.", "success")
    return redirect(url_for("reunioes.detalhes_reuniao", id=reuniao.id))


def limpar_cabecalho_ata(texto):
    linhas = texto.splitlines()
    resultado = []
    pulando = True

    for linha in linhas:
        l = linha.strip()

        if pulando:
            if l.startswith("Aos ") or l.startswith("Ao "):
                pulando = False
                resultado.append(l)
            continue

        resultado.append(linha)

    return "\n".join(resultado)


@reunioes.route("/<int:id>/ata/visualizar")
@login_required
@administrativo_required
def visualizar_ata(id):
    reuniao = Reuniao.query.get_or_404(id)

    texto_ata = limpar_cabecalho_ata(reuniao.ata or "")

    texto_ata = " ".join(
        linha.strip()
        for linha in texto_ata.splitlines()
        if linha.strip()
    )

    ata_linhas = textwrap.wrap(
        texto_ata,
        width=88,
        break_long_words=False,
        break_on_hyphens=False
    )

    return render_template(
        "documentos/ata_modelo.html",
        reuniao=reuniao,
        ata_linhas=ata_linhas
    )

@reunioes.route("/<int:id>/ata/word")
@login_required
@administrativo_required
def baixar_ata_word(id):
    from io import BytesIO
    from flask import send_file
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    reuniao = Reuniao.query.get_or_404(id)

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Logo central
    logo = doc.add_paragraph()
    logo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        run_logo = logo.add_run()
        run_logo.add_picture("app/static/img/ppeb.png", width=Inches(0.85))
    except Exception:
        pass

    # Cabeçalho institucional
    cabecalho = doc.add_paragraph()
    cabecalho.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cabecalho.add_run(
        "SERVIÇO PÚBLICO FEDERAL\n"
        "UNIVERSIDADE FEDERAL DO PARÁ\n"
        "NÚCLEO DE ESTUDOS TRANSDISCIPLINARES EM EDUCAÇÃO BÁSICA\n"
        "PROGRAMA DE PÓS-GRADUAÇÃO EM CURRÍCULO E GESTÃO DA ESCOLA BÁSICA"
    )
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)

    # Título da ata
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    titulo.paragraph_format.left_indent = Inches(3.1)

    run_titulo = titulo.add_run(
        f"Ata da {reuniao.titulo} do Colegiado do\n"
        "Programa de Pós-Graduação em Currículo e\n"
        "Gestão da Escola Básica do NEB/UFPA,\n"
        f"realizada no dia {reuniao.data.strftime('%d/%m/%Y')}"
    )
    run_titulo.bold = True
    run_titulo.font.name = "Times New Roman"
    run_titulo.font.size = Pt(12)

    doc.add_paragraph("")

    # Corpo da ata sem repetir cabeçalho
    texto_ata = limpar_cabecalho_ata(reuniao.ata or "")

    if not texto_ata.strip():
        texto_ata = "Ata não cadastrada."

    paragrafos = [
        linha.strip()
        for linha in texto_ata.splitlines()
        if linha.strip()
    ]

    if not paragrafos:
        paragrafos = [texto_ata.strip()]

    for paragrafo in paragrafos:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Inches(0.35)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)

        run = p.add_run(paragrafo)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)

    arquivo = BytesIO()
    doc.save(arquivo)
    arquivo.seek(0)

    return send_file(
        arquivo,
        as_attachment=True,
        download_name=f"ata_reuniao_{reuniao.id}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
@reunioes.route("/<int:id>/encerrar-justificativas", methods=["POST"])
@login_required
@administrativo_required
def encerrar_justificativas(id):
    reuniao = Reuniao.query.get_or_404(id)

    reuniao.justificativas_encerradas = True
    db.session.commit()

    flash("Justificativas encerradas com sucesso.", "success")
    return redirect(url_for("reunioes.detalhes_reuniao", id=reuniao.id))


@reunioes.route("/<int:id>/reabrir-justificativas", methods=["POST"])
@login_required
@administrativo_required
def reabrir_justificativas(id):
    reuniao = Reuniao.query.get_or_404(id)

    reuniao.justificativas_encerradas = False
    db.session.commit()

    flash("Justificativas reabertas com sucesso.", "success")
    return redirect(url_for("reunioes.detalhes_reuniao", id=reuniao.id))


@reunioes.route("/<int:id>/enviar-pauta", methods=["POST"])
@login_required
@administrativo_required
def enviar_pauta_por_email(id):
    reuniao = Reuniao.query.get_or_404(id)

    frequencias = Frequencia.query.filter_by(
        reuniao_id=reuniao.id
    ).all()

    emails = [
        frequencia.membro.email
        for frequencia in frequencias
        if frequencia.membro
        and frequencia.membro.ativo
        and frequencia.membro.email
    ]

    if not emails:
        flash("Nenhum e-mail cadastrado para os membros desta reunião.", "warning")
        return redirect(url_for("reunioes.detalhes_reuniao", id=reuniao.id))

    link_pauta = url_for(
        "reunioes.visualizar_pauta",
        id=reuniao.id,
        _external=True
    )

    mensagem = Message(
        subject=f"Pauta da reunião: {reuniao.titulo}",
        sender=current_app.config["MAIL_USERNAME"],
        recipients=[current_app.config["MAIL_USERNAME"]],
        bcc=emails,
        body=f"""
Prezados(as),

A pauta da reunião "{reuniao.titulo}" está disponível para visualização.

Acesse pelo link abaixo:
{link_pauta}

Atenciosamente,
Coordenação do PPEB
"""
    )

    mail.send(mensagem)

    flash("Link da pauta enviado para os membros da reunião.", "success")
    return redirect(url_for("reunioes.detalhes_reuniao", id=reuniao.id))

@reunioes.route('/excluir/<int:id>', methods=['POST'])
@login_required
@administrativo_required
def excluir_reuniao(id):
    reuniao = Reuniao.query.get_or_404(id)

    Frequencia.query.filter_by(reuniao_id=reuniao.id).delete()

    db.session.delete(reuniao)
    db.session.commit()

    flash("Reunião excluída com sucesso.", "success")
    return redirect(url_for('reunioes.listar_reunioes'))

@reunioes.route("/pauta/<int:id>/visualizar")
def visualizar_pauta(id):
    reuniao = Reuniao.query.get_or_404(id)
    return render_template(
        "documentos/visualizar_pauta.html",
        reuniao=reuniao
    )

@reunioes.route("/<int:reuniao_id>/frequencia/<int:frequencia_id>/remover", methods=["POST"])
@login_required
@administrativo_required
def remover_frequencia_reuniao(reuniao_id, frequencia_id):
    frequencia = Frequencia.query.get_or_404(frequencia_id)

    if frequencia.reuniao_id != reuniao_id:
        flash("Frequência inválida para esta reunião.", "error")
        return redirect(url_for("reunioes.detalhes_reuniao", id=reuniao_id))

    frequencia.removido = True

    db.session.commit()

    flash("Participante removido desta reunião.", "success")
    return redirect(url_for("reunioes.detalhes_reuniao", id=reuniao_id))